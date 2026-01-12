import io
import json
import os
import re
import time
import uuid
from pathlib import Path

import boto3
from docx import Document
from openai import OpenAI, APIError, APIConnectionError, RateLimitError, APITimeoutError
from pgvector.sqlalchemy import Vector
from pypdf import PdfReader
import yaml

from negotiation_templates import get_template
from celery import Celery
from sqlalchemy import Column, DateTime, Integer, String, Text, create_engine, delete, func, select
from sqlalchemy.dialects.postgresql import JSONB, UUID
from sqlalchemy.orm import Session, declarative_base, sessionmaker

redis_url = os.getenv("DPA_REDIS_URL", "redis://localhost:6379/0")
database_url = os.getenv("DPA_DATABASE_URL", "postgresql+psycopg2://dpa:dpa@localhost:5432/dpa")
s3_endpoint = os.getenv("S3_INTERNAL_ENDPOINT", os.getenv("DPA_S3_ENDPOINT", "http://localhost:9000"))
s3_access_key = os.getenv("DPA_S3_ACCESS_KEY", "minio")
s3_secret_key = os.getenv("DPA_S3_SECRET_KEY", "minio123")
s3_bucket = os.getenv("DPA_S3_BUCKET", "dpa-guard")
s3_region = os.getenv("DPA_S3_REGION", "us-east-1")
rag_enabled = os.getenv("RAG_ENABLED", "false").lower() in {"1", "true", "yes"}
openai_api_key = os.getenv("OPENAI_API_KEY")
openai_embed_model = os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small")

celery_app = Celery("dpa_guard", broker=redis_url, backend=redis_url)

engine = create_engine(database_url, pool_pre_ping=True)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

Base = declarative_base()

PLAYBOOK_DIR = Path(__file__).parent / "playbooks"

NON_DPA_KEYWORDS = {"prd", "tdd", "use case", "requirements", "mvp", "roadmap"}
DPA_KEYWORDS = {
    "data processing agreement",
    "data processing addendum",
    "controller",
    "processor",
    "article 28",
    "gdpr",
    "annex",
    "schedule",
}
PRD_TDD_KEYWORDS = {
    "product requirements document",
    "technical design document",
    "mvp",
    "use cases",
    "non-functional requirements",
    "last updated",
}


class Review(Base):
    __tablename__ = "reviews"

    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    status = Column(String(32), nullable=False, default="draft")
    context_json = Column(JSONB, nullable=True)
    results_json = Column(JSONB, nullable=True)
    extracted_text = Column(Text, nullable=True)
    extracted_meta = Column(JSONB, nullable=True)
    playbook_id = Column(UUID(as_uuid=True), nullable=True)
    source_object_key = Column(Text, nullable=True)
    source_filename = Column(Text, nullable=True)
    source_mime = Column(Text, nullable=True)
    error_message = Column(Text, nullable=True)
    created_at = Column(DateTime(timezone=True), server_default=func.now(), nullable=False)
    updated_at = Column(
        DateTime(timezone=True),
        server_default=func.now(),
        onupdate=func.now(),
        nullable=False,
    )


class Playbook(Base):
    __tablename__ = "playbooks"

    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    title = Column(Text, nullable=True)
    version = Column(String(32), nullable=True)
    source_object_key = Column(Text, nullable=True)
    source_filename = Column(Text, nullable=True)
    status = Column(String(32), nullable=False, default="uploaded")
    error_message = Column(Text, nullable=True)
    created_at = Column(DateTime(timezone=True), server_default=func.now(), nullable=False)
    updated_at = Column(
        DateTime(timezone=True),
        server_default=func.now(),
        onupdate=func.now(),
        nullable=False,
    )


class PlaybookChunk(Base):
    __tablename__ = "playbook_chunks"

    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    playbook_id = Column(UUID(as_uuid=True), nullable=False, index=True)
    chunk_index = Column(Integer, nullable=False)
    content = Column(Text, nullable=False)
    meta_json = Column(JSONB, nullable=True)
    embedding = Column(Vector(1536), nullable=True)
    created_at = Column(DateTime(timezone=True), server_default=func.now(), nullable=False)


def get_s3_client():
    return boto3.client(
        "s3",
        endpoint_url=s3_endpoint,
        aws_access_key_id=s3_access_key,
        aws_secret_access_key=s3_secret_key,
        region_name=s3_region,
    )


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    collapsed = []
    blank_streak = 0
    for line in lines:
        if line:
            if blank_streak:
                collapsed.append("")
                blank_streak = 0
            collapsed.append(" ".join(line.split()))
        else:
            blank_streak += 1
    return "\n".join(collapsed).strip()


def chunk_text(
    text: str, chunk_size: int = 2000, overlap: int = 200
) -> list[dict[str, int]]:
    chunks: list[dict[str, int]] = []
    if not text:
        return chunks
    step = max(chunk_size - overlap, 1)
    start = 0
    index = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        chunks.append({"index": index, "start_char": start, "end_char": end})
        index += 1
        start += step
    return chunks


def _retry(func, attempts: int = 2):
    for attempt in range(attempts):
        try:
            return func()
        except (APIConnectionError, APIError, RateLimitError, APITimeoutError):
            if attempt + 1 >= attempts:
                raise
            time.sleep(0.5)


def embed_texts_openai(texts: list[str]) -> list[list[float]]:
    if not rag_enabled:
        raise RuntimeError("RAG is disabled")
    if not openai_api_key:
        raise RuntimeError("OPENAI_API_KEY is required for RAG indexing")
    client = OpenAI(api_key=openai_api_key)

    def _call():
        response = client.embeddings.create(
            model=openai_embed_model,
            input=texts,
            timeout=30,
        )
        return [item.embedding for item in response.data]

    return _retry(_call)


def generate_structured_openai(prompt: str, json_schema: dict) -> tuple[str, dict | None]:
    if not llm_enabled or llm_provider != "openai":
        raise RuntimeError("LLM is disabled")
    if not openai_api_key:
        raise RuntimeError("OPENAI_API_KEY is required for LLM generation")
    client = OpenAI(api_key=openai_api_key)

    def _call():
        response = client.chat.completions.create(
            model=openai_model,
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_schema", "json_schema": json_schema},
            timeout=30,
        )
        content = response.choices[0].message.content or ""
        parsed = None
        try:
            parsed = json.loads(content)
        except json.JSONDecodeError:
            start = content.find("{")
            end = content.rfind("}")
            if start != -1 and end != -1 and end > start:
                try:
                    parsed = json.loads(content[start : end + 1])
                except json.JSONDecodeError:
                    parsed = None
        return content, parsed

    return _retry(_call)


def _truncate(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + "…"


def build_llm_prompt(context: dict, risk_table: list[dict]) -> str:
    findings = []
    for finding in risk_table:
        rag_chunks = []
        rag = finding.get("rag") or {}
        for chunk in rag.get("chunks", [])[:2]:
            rag_chunks.append(
                {
                    "chunk_index": chunk.get("chunk_index"),
                    "content": _truncate(chunk.get("content", ""), 600),
                    "meta_json": chunk.get("meta_json"),
                }
            )
        findings.append(
            {
                "check_id": finding.get("check_id"),
                "title": finding.get("title"),
                "severity": finding.get("severity"),
                "what_good_looks_like": finding.get("what_good_looks_like"),
                "recommendation": finding.get("recommendation"),
                "evidence_quotes": [
                    {
                        "quote": quote.get("quote"),
                        "chunk_index": quote.get("chunk_index"),
                        "start_char": quote.get("start_char"),
                        "end_char": quote.get("end_char"),
                    }
                    for quote in (finding.get("evidence_quotes") or [])
                ],
                "rag_chunks": rag_chunks,
            }
        )
    payload = {
        "context": context or {},
        "risk_table": findings,
    }
    return (
        "You are generating an executive summary and negotiation pack based only on the provided "
        "risk table findings and evidence quotes. Do not create new evidence. Only include check_ids "
        "present in the risk_table.\n\n"
        f"INPUT JSON:\n{json.dumps(payload, ensure_ascii=False)}"
    )


def _trim_text(value: str | None, limit: int) -> str:
    if not value:
        return ""
    if len(value) <= limit:
        return value
    return value[:limit]


def build_deterministic_pack(findings: list[dict]) -> list[dict]:
    pack = []
    for finding in findings:
        template = get_template(finding.get("check_id", ""))
        pack.append(
            {
                "check_id": finding.get("check_id"),
                "title": finding.get("title"),
                "severity": finding.get("severity"),
                "ask": template.get("ask"),
                "fallback": template.get("fallback"),
                "rationale": template.get("rationale"),
            }
        )
    return pack


def clean_llm_output(
    output: dict,
    allowed_check_ids: set[str],
    deterministic_exec_summary: str,
    deterministic_pack: list[dict],
    max_len: int = 400,
) -> dict:
    exec_summary = output.get("exec_summary", "")
    if not isinstance(exec_summary, str) or not exec_summary.strip():
        exec_summary = deterministic_exec_summary
    pack_from_llm = output.get("negotiation_pack", [])
    deterministic_map = {item.get("check_id"): item for item in deterministic_pack}
    cleaned_pack = []
    seen = set()

    if isinstance(pack_from_llm, list):
        for item in pack_from_llm:
            if not isinstance(item, dict):
                continue
            check_id = item.get("check_id")
            if check_id not in allowed_check_ids or check_id in seen:
                continue
            base = deterministic_map.get(check_id, {})
            cleaned_pack.append(
                {
                    "check_id": check_id,
                    "title": base.get("title"),
                    "severity": base.get("severity"),
                    "ask": _trim_text(item.get("ask") or base.get("ask"), max_len),
                    "fallback": _trim_text(item.get("fallback") or base.get("fallback"), max_len),
                    "rationale": _trim_text(item.get("rationale") or base.get("rationale"), max_len),
                }
            )
            seen.add(check_id)

    for item in deterministic_pack:
        check_id = item.get("check_id")
        if check_id in allowed_check_ids and check_id not in seen:
            cleaned_pack.append(
                {
                    "check_id": check_id,
                    "title": item.get("title"),
                    "severity": item.get("severity"),
                    "ask": _trim_text(item.get("ask"), max_len),
                    "fallback": _trim_text(item.get("fallback"), max_len),
                    "rationale": _trim_text(item.get("rationale"), max_len),
                }
            )
            seen.add(check_id)

    return {"exec_summary": exec_summary, "negotiation_pack": cleaned_pack}


def maybe_apply_llm(
    review: Review, results: dict, context_json: dict | None
) -> dict:
    if not (llm_enabled and llm_provider == "openai" and openai_api_key):
        return results
    base_results = results if isinstance(results, dict) else {}
    results_copy = dict(base_results)
    risk_table = results_copy.get("risk_table") or []
    if not risk_table:
        return results_copy
    allowed_check_ids = {finding.get("check_id") for finding in risk_table if finding.get("check_id")}
    prompt = build_llm_prompt(context_json or {}, risk_table)
    schema = {
        "name": "exec_and_negotiation",
        "schema": {
            "type": "object",
            "properties": {
                "exec_summary": {"type": "string"},
                "negotiation_pack": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "check_id": {"type": "string"},
                            "ask": {"type": "string"},
                            "fallback": {"type": "string"},
                            "rationale": {"type": "string"},
                        },
                        "required": ["check_id", "ask", "fallback", "rationale"],
                    },
                },
            },
            "required": ["exec_summary", "negotiation_pack"],
            "additionalProperties": False,
        },
    }
    deterministic_exec_summary = summarize_findings(risk_table)
    deterministic_pack = build_deterministic_pack(risk_table)
    try:
        raw_text, parsed = generate_structured_openai(prompt, schema)
    except Exception:  # noqa: BLE001
        results_copy["exec_summary"] = deterministic_exec_summary
        results_copy["negotiation_pack"] = deterministic_pack
        return results_copy
    if not isinstance(parsed, dict):
        cleaned = {
            "exec_summary": deterministic_exec_summary,
            "negotiation_pack": deterministic_pack,
        }
    else:
        cleaned = clean_llm_output(parsed, allowed_check_ids, deterministic_exec_summary, deterministic_pack)
    generated_at = datetime.now(timezone.utc).isoformat()
    results_copy["llm"] = {
        "model": openai_model,
        "generated_at": generated_at,
        "raw": raw_text,
        "cleaned": cleaned,
    }
    results_copy["exec_summary"] = cleaned["exec_summary"]
    results_copy["negotiation_pack"] = cleaned["negotiation_pack"]
    return results_copy
def retrieve_playbook_chunks(
    db: Session, playbook_id: uuid.UUID, query_text: str, k: int = 6
) -> list[PlaybookChunk]:
    embedding = embed_texts_openai([query_text])[0]
    stmt = (
        select(PlaybookChunk)
        .where(PlaybookChunk.playbook_id == playbook_id)
        .order_by(PlaybookChunk.embedding.cosine_distance(embedding))
        .limit(k)
    )
    return db.scalars(stmt).all()


def _extract_playbook_text(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    if lower.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        return "\n\n".join(p.text for p in doc.paragraphs)
    return data.decode("utf-8", errors="ignore")


def _parse_structured_playbook(data: bytes) -> list[dict] | None:
    text = data.decode("utf-8", errors="ignore")
    try:
        parsed = yaml.safe_load(text)
    except Exception:
        parsed = None
    if not parsed:
        try:
            parsed = json.loads(text)
        except Exception:
            parsed = None
    if not isinstance(parsed, dict) or "checks" not in parsed:
        return None
    checks = parsed.get("checks") or []
    if not isinstance(checks, list):
        return None
    chunks = []
    for idx, check in enumerate(checks):
        if not isinstance(check, dict):
            continue
        text_parts = [
            check.get("title"),
            check.get("what_good_looks_like"),
            check.get("recommendation"),
            " ".join(check.get("patterns") or []),
        ]
        content = "\n".join(part for part in text_parts if part)
        if not content:
            continue
        chunks.append(
            {
                "chunk_index": idx,
                "content": content,
                "meta_json": {
                    "check_id": check.get("check_id"),
                    "severity": check.get("severity"),
                },
            }
        )
    return chunks


def validate_quote(extracted_text: str, quote: str, start: int, end: int) -> tuple[int, int] | None:
    if start >= 0 and end <= len(extracted_text) and extracted_text[start:end] == quote:
        return start, end
    fallback = extracted_text.find(quote)
    if fallback != -1:
        return fallback, fallback + len(quote)
    return None


SECTION_HEADING_RE = re.compile(r"^(?P<heading>[A-Z]{2,10}-[A-Z]{2,10}-\d{2}\b.*)$", re.MULTILINE)


def _split_playbook_sections(text: str, chunk_size: int = 1800, overlap: int = 200) -> list[dict]:
    matches = list(SECTION_HEADING_RE.finditer(text))
    if not matches:
        return []
    sections = []
    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        heading = match.group("heading").strip()
        sections.append({"heading": heading, "start": start, "end": end})

    chunks = []
    chunk_index = 0
    for section in sections:
        section_text = text[section["start"] : section["end"]]
        if len(section_text) <= chunk_size + overlap:
            chunks.append(
                {
                    "chunk_index": chunk_index,
                    "content": section_text,
                    "meta_json": {
                        "heading": section["heading"],
                        "start_char": section["start"],
                        "end_char": section["end"],
                    },
                }
            )
            chunk_index += 1
            continue

        sub_ranges = chunk_text(section_text, chunk_size=chunk_size, overlap=overlap)
        for sub in sub_ranges:
            sub_start = section["start"] + sub["start_char"]
            sub_end = section["start"] + sub["end_char"]
            content = section_text[sub["start_char"] : sub["end_char"]]
            if sub["start_char"] != 0:
                content = f"{section['heading']}\n{content}"
            chunks.append(
                {
                    "chunk_index": chunk_index,
                    "content": content,
                    "meta_json": {
                        "heading": section["heading"],
                        "start_char": sub_start,
                        "end_char": sub_end,
                    },
                }
            )
            chunk_index += 1
    return chunks


def _tokenize_title(title: str | None) -> list[str]:
    if not title:
        return []
    return [token for token in re.split(r"[^a-zA-Z0-9]+", title.lower()) if len(token) >= 4]


def _get_chunk_heading(chunk) -> str | None:
    meta = None
    if isinstance(chunk, dict):
        meta = chunk.get("meta_json")
    else:
        meta = getattr(chunk, "meta_json", None)
    if isinstance(meta, dict):
        heading = meta.get("heading")
        if isinstance(heading, str) and heading.strip():
            return heading
    return None


def _score_chunk_for_finding(finding: dict, heading: str | None) -> int:
    if not heading:
        return 0
    heading_lower = heading.lower()
    score = 0
    title_tokens = _tokenize_title(finding.get("title"))
    if any(token in heading_lower for token in title_tokens):
        score += 3
    check_id = finding.get("check_id") or ""
    prefix_map = {
        "breach_notification_timeline": ["BR"],
        "subprocessor_authorization": ["SUB"],
        "international_transfers": ["TR", "TRANSFER"],
        "audit_rights": ["AUDIT"],
        "deletion_or_return": ["DEL"],
        "confidentiality_of_personnel": ["CONF"],
        "purpose_limitation_and_instructions": ["INSTR", "PURPOSE"],
    }
    for prefix in prefix_map.get(check_id, []):
        if prefix.lower() in heading_lower:
            score += 2
            break
    return score


def select_rag_chunks_for_finding(chunks: list, finding: dict, max_chunks: int = 2) -> list:
    scored = []
    for idx, chunk in enumerate(chunks):
        heading = _get_chunk_heading(chunk)
        score = _score_chunk_for_finding(finding, heading)
        scored.append((score, idx, chunk))
    scored.sort(key=lambda item: (-item[0], item[1]))
    if not scored:
        return []
    best_score, _, best_chunk = scored[0]
    selected = [best_chunk]
    if max_chunks > 1 and len(scored) > 1:
        second_score, _, second_chunk = scored[1]
        if second_score > 0 and second_score == best_score:
            selected.append(second_chunk)
    return selected


def load_playbook(playbook_id: str = "eu_controller") -> dict:
    for path in PLAYBOOK_DIR.glob("*.yaml"):
        data = yaml.safe_load(path.read_text())
        if data.get("id") == playbook_id:
            return data
    raise ValueError(f"Playbook not found: {playbook_id}")


def select_playbook_id(context_json: dict | None) -> str:
    if not context_json:
        return "eu_controller"
    context = context_json.get("context") if isinstance(context_json, dict) else None
    if isinstance(context, dict):
        region = context.get("region")
        role = context.get("company_role")
    else:
        region = context_json.get("region")
        role = context_json.get("company_role")
    if (region or "").lower() in {"eu", "europe", "european_union"}:
        return "eu_controller"
    if (role or "").lower() == "controller":
        return "eu_controller"
    return "eu_controller"


def find_chunk_index(chunks: list[dict[str, int]], start_char: int, end_char: int) -> int | None:
    for chunk in chunks:
        if chunk["start_char"] <= start_char < chunk["end_char"]:
            return chunk["index"]
    return None


def run_checks(extracted_text: str, chunks: list[dict[str, int]], playbook: dict) -> list[dict]:
    findings: list[dict] = []
    lowered = extracted_text.lower()
    for check in playbook.get("checks", []):
        evidence_quotes: list[dict] = []
        for pattern in check.get("patterns", []):
            for match in re.finditer(pattern, extracted_text, flags=re.IGNORECASE | re.MULTILINE):
                if len(evidence_quotes) >= 3:
                    break
                start = match.start()
                end = match.end()
                snippet_start = max(0, start - 80)
                snippet_end = min(len(extracted_text), snippet_start + 240)
                quote = extracted_text[snippet_start:snippet_end]
                if any(keyword in quote.lower() for keyword in NON_DPA_KEYWORDS):
                    continue
                validated = validate_quote(extracted_text, quote, snippet_start, snippet_end)
                if not validated:
                    continue
                v_start, v_end = validated
                chunk_index = find_chunk_index(chunks, v_start, v_end)
                evidence_quotes.append(
                    {
                        "quote": extracted_text[v_start:v_end],
                        "start_char": v_start,
                        "end_char": v_end,
                        "chunk_index": chunk_index,
                    }
                )
            if len(evidence_quotes) >= 3:
                break
        if not evidence_quotes:
            continue
        findings.append(
            {
                "check_id": check.get("check_id"),
                "title": check.get("title"),
                "severity": check.get("severity"),
                "what_good_looks_like": check.get("what_good_looks_like"),
                "recommendation": check.get("recommendation"),
                "evidence_quotes": evidence_quotes,
            }
        )
    return findings


def summarize_findings(findings: list[dict]) -> str:
    counts = {"high": 0, "medium": 0, "med": 0, "low": 0}
    for finding in findings:
        severity = (finding.get("severity") or "").lower()
        if severity in counts:
            counts[severity] += 1
    total_high = counts["high"]
    total_med = counts["medium"] + counts["med"]
    total_low = counts["low"]
    top_titles = [f["title"] for f in findings[:3] if f.get("title")]
    top_part = ", ".join(top_titles) if top_titles else "none"
    return f"Findings: high={total_high}, med={total_med}, low={total_low}. Top: {top_part}."


def detect_doc_type(extracted_text: str) -> dict:
    lowered = extracted_text.lower()
    dpa_score = sum(1 for keyword in DPA_KEYWORDS if keyword in lowered)
    non_dpa_score = sum(1 for keyword in PRD_TDD_KEYWORDS if keyword in lowered)
    warnings = []
    if non_dpa_score >= max(2, dpa_score + 1):
        warnings.append("Document appears non-DPA (PRD/TDD indicators detected); results may be unreliable.")
    return {"dpa": dpa_score, "non_dpa": non_dpa_score, "warnings": warnings}


@celery_app.task(name="process_review")
def process_review(review_id: str):
    # IMPORTANT: never log raw contract text; only IDs/status/metadata.
    db: Session = SessionLocal()
    try:
        try:
            review_uuid = uuid.UUID(review_id)
        except ValueError:
            return {"status": "failed"}

        review = db.scalar(select(Review).where(Review.id == review_uuid))
        if not review:
            return {"status": "not_found"}
        if not review.source_object_key:
            review.status = "failed"
            review.error_message = "Source file missing"
            db.add(review)
            db.commit()
            return {"status": "failed"}

        s3 = get_s3_client()
        head = s3.head_object(Bucket=s3_bucket, Key=review.source_object_key)
        size = head.get("ContentLength", 0)
        if not size or size <= 0:
            raise ValueError("Source file empty")

        obj = s3.get_object(Bucket=s3_bucket, Key=review.source_object_key)
        data = obj["Body"].read()
        if not data:
            raise ValueError("Source file empty")

        extracted_text = ""
        pages = None
        if review.source_mime == "application/pdf":
            reader = PdfReader(io.BytesIO(data))
            pages = len(reader.pages)
            extracted_text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
        elif review.source_mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(io.BytesIO(data))
            extracted_text = "\n\n".join(p.text for p in doc.paragraphs)
        else:
            raise ValueError("Unsupported source type")

        normalized = normalize_text(extracted_text)
        if review.source_mime == "application/pdf" and len(normalized) < 500:
            review.status = "failed"
            review.error_message = (
                "PDF appears scanned or contains insufficient text; OCR not supported in MVP."
            )
            review.extracted_text = normalized
            review.extracted_meta = {"chars": len(normalized), "pages": pages, "chunks": []}
            db.add(review)
            db.commit()
            return {"status": "failed"}

        chunks = chunk_text(normalized)
        review.extracted_text = normalized
        review.extracted_meta = {
            "chars": len(normalized),
            "pages": pages,
            "chunks": chunks,
        }

        playbook_id = select_playbook_id(review.context_json or {})
        playbook = load_playbook(playbook_id)
        findings = run_checks(normalized, chunks, playbook)
        exec_summary = summarize_findings(findings)
        doc_type = detect_doc_type(normalized)
        if doc_type["warnings"]:
            exec_summary = f"{doc_type['warnings'][0]} {exec_summary}"

        rag_payload = None
        per_finding_rag = {}
        if rag_enabled:
            selected_playbook_id = review.playbook_id
            if not selected_playbook_id:
                selected_playbook = db.scalar(
                    select(Playbook).where(Playbook.status == "ready").order_by(Playbook.created_at.desc())
                )
                selected_playbook_id = selected_playbook.id if selected_playbook else None
            if selected_playbook_id:
                queries = [
                    "breach notification timeline",
                    "subprocessor objection and authorization",
                    "international transfers and SCCs",
                    "audit and inspection rights",
                    "deletion or return of data upon termination",
                    "processing on documented instructions",
                ]
                top_chunks = []
                for query in queries:
                    chunks_found = retrieve_playbook_chunks(db, selected_playbook_id, query, k=6)
                    top_chunks.append(
                        {
                            "query": query,
                            "chunks": [
                                {
                                    "chunk_index": chunk.chunk_index,
                                    "content": chunk.content,
                                    "meta_json": chunk.meta_json,
                                }
                                for chunk in chunks_found
                            ],
                        }
                    )
                rag_payload = {
                    "playbook_id": str(selected_playbook_id),
                    "queries": queries,
                    "top_chunks": top_chunks,
                }
                per_finding_rag = {"playbook_id": str(selected_playbook_id)}

        if rag_enabled and per_finding_rag.get("playbook_id"):
            max_total_chars = 1200
            for finding in findings:
                query = f"{finding.get('check_id')} {finding.get('title')} {finding.get('recommendation') or ''}"
                chunks_found = retrieve_playbook_chunks(
                    db, uuid.UUID(per_finding_rag["playbook_id"]), query, k=6
                )
                selected_chunks = select_rag_chunks_for_finding(chunks_found, finding, max_chunks=2)
                rag_chunks = []
                total_chars = 0
                for chunk in selected_chunks:
                    content = _truncate(chunk.content or "", 600)
                    if not content:
                        continue
                    remaining = max_total_chars - total_chars
                    if remaining <= 0:
                        break
                    if len(content) > remaining:
                        content = content[:remaining]
                    rag_chunks.append(
                        {
                            "chunk_index": chunk.chunk_index,
                            "content": content,
                            "meta_json": chunk.meta_json,
                        }
                    )
                    total_chars += len(content)
                finding["rag"] = {"chunks": rag_chunks}

        negotiation_pack = build_deterministic_pack(findings)

        review.results_json = {
            "playbook": {
                "id": playbook.get("id"),
                "version": playbook.get("version"),
                "title": playbook.get("title"),
            },
            "exec_summary": exec_summary,
            "risk_table": findings,
            "negotiation_pack": negotiation_pack,
            "extraction": {
                "chars": len(normalized),
                "pages": pages,
                "chunks": len(chunks),
                "doc_type_score": {"dpa": doc_type["dpa"], "non_dpa": doc_type["non_dpa"]},
            },
        }
        if rag_payload:
            review.results_json["rag"] = rag_payload
        review.results_json = maybe_apply_llm(review, review.results_json, review.context_json)
        if doc_type["warnings"]:
            review.results_json["warnings"] = doc_type["warnings"]
        review.status = "completed"
        review.error_message = None
        db.add(review)
        db.commit()
        return {"status": "completed"}
    except Exception as exc:  # noqa: BLE001
        review = db.scalar(select(Review).where(Review.id == review_uuid))
        if review:
            review.status = "failed"
            review.error_message = str(exc)
            db.add(review)
            db.commit()
        return {"status": "failed"}
    finally:
        db.close()


@celery_app.task(name="index_playbook")
def index_playbook(playbook_id: str):
    # IMPORTANT: never log raw playbook text; only IDs/status/metadata.
    db: Session = SessionLocal()
    try:
        try:
            playbook_uuid = uuid.UUID(playbook_id)
        except ValueError:
            return {"status": "failed"}

        playbook = db.scalar(select(Playbook).where(Playbook.id == playbook_uuid))
        if not playbook:
            return {"status": "not_found"}

        playbook.status = "indexing"
        playbook.error_message = None
        db.add(playbook)
        db.commit()

        if not rag_enabled:
            playbook.status = "failed"
            playbook.error_message = "RAG is disabled"
            db.add(playbook)
            db.commit()
            return {"status": "failed"}

        s3 = get_s3_client()
        obj = s3.get_object(Bucket=s3_bucket, Key=playbook.source_object_key)
        data = obj["Body"].read()

        structured_chunks = _parse_structured_playbook(data)
        chunks = []
        if structured_chunks is not None:
            chunks = structured_chunks
        else:
            text = normalize_text(_extract_playbook_text(playbook.source_filename or "", data))
            sections = _split_playbook_sections(text)
            if sections:
                chunks = sections
            else:
                chunk_ranges = chunk_text(text, chunk_size=1200, overlap=150)
                for chunk in chunk_ranges:
                    content = text[chunk["start_char"] : chunk["end_char"]]
                    chunks.append(
                        {
                            "chunk_index": chunk["index"],
                            "content": content,
                            "meta_json": {"start_char": chunk["start_char"], "end_char": chunk["end_char"]},
                        }
                    )

        embeddings = embed_texts_openai([chunk["content"] for chunk in chunks])
        for chunk, embedding in zip(chunks, embeddings, strict=False):
            db.add(
                PlaybookChunk(
                    playbook_id=playbook_uuid,
                    chunk_index=chunk["chunk_index"],
                    content=chunk["content"],
                    meta_json=chunk.get("meta_json"),
                    embedding=embedding,
                )
            )
        playbook.status = "ready"
        db.add(playbook)
        db.commit()
        return {"status": "ready"}
    except Exception as exc:  # noqa: BLE001
        playbook = None
        if "playbook_uuid" in locals():
            playbook = db.scalar(select(Playbook).where(Playbook.id == playbook_uuid))
        if playbook:
            playbook.status = "failed"
            playbook.error_message = str(exc)
            db.add(playbook)
            db.commit()
        return {"status": "failed"}
    finally:
        db.close()


@celery_app.task(name="reindex_playbook")
def reindex_playbook(playbook_id: str):
    db: Session = SessionLocal()
    try:
        try:
            playbook_uuid = uuid.UUID(playbook_id)
        except ValueError:
            return {"status": "failed"}
        db.execute(delete(PlaybookChunk).where(PlaybookChunk.playbook_id == playbook_uuid))
        db.commit()
    finally:
        db.close()
    return index_playbook(playbook_id)


@celery_app.task(name="rerun_llm")
def rerun_llm(review_id: str):
    db: Session = SessionLocal()
    try:
        try:
            review_uuid = uuid.UUID(review_id)
        except ValueError:
            return {"status": "failed"}
        review = db.scalar(select(Review).where(Review.id == review_uuid))
        if not review or not review.results_json:
            return {"status": "not_found"}
        review.results_json = maybe_apply_llm(review, review.results_json, review.context_json)
        db.add(review)
        db.commit()
        return {"status": "completed"}
    finally:
        db.close()
from datetime import datetime, timezone
llm_enabled = os.getenv("LLM_ENABLED", "false").lower() in {"1", "true", "yes"}
llm_provider = os.getenv("LLM_PROVIDER", "openai").lower()
openai_model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
